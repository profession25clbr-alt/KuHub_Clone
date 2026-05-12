"""
Script para generar EVALUACION_PARCIAL_2_TPY1101.docx
Incluye todo el contenido de DIAGRAMA_ER_KUHUB.md y ARQUITECTURA_4+1_VIEWS.md
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


def set_cell_shading(cell, color):
    """Aplica color de fondo a una celda."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    shading_elm.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_table_from_data(doc, headers, rows, col_widths=None):
    """Crea una tabla formateada en el documento."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        set_cell_shading(cell, '2F5496')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            if r_idx % 2 == 0:
                set_cell_shading(cell, 'D9E2F3')

    return table


def add_code_block(doc, text, title=None):
    """Agrega un bloque de código/diagrama con formato monoespaciado."""
    if title:
        doc.add_paragraph(title, style='Heading 4')
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.5)
    run = para.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    return para


def create_document():
    doc = Document()

    # ══════════════════════════════════════════════════════════════
    # CONFIGURACIÓN DE ESTILOS
    # ══════════════════════════════════════════════════════════════
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ══════════════════════════════════════════════════════════════
    # PORTADA
    # ══════════════════════════════════════════════════════════════
    for _ in range(4):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('EVALUACIÓN PARCIAL 2')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(47, 84, 150)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('TPY1101 — Taller de Proyecto')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(47, 84, 150)

    doc.add_paragraph()

    project = doc.add_paragraph()
    project.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = project.add_run('KuHub — Sistema de Gestión de Bodega e Inventario')
    run.font.size = Pt(16)
    run.bold = True

    doc.add_paragraph()

    version = doc.add_paragraph()
    version.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = version.add_run('Versión del Sistema: v1.0.8')
    run.font.size = Pt(12)

    doc.add_paragraph()

    info_lines = [
        'Base de Datos: PostgreSQL 16.13',
        'Backend: Spring Boot 3 + Java 21',
        'Frontend: React 18 + TypeScript + Vite',
        'Infraestructura: AWS Lightsail',
        'Fecha: 12 de mayo de 2026',
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(line)
        run.font.size = Pt(11)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════
    # ÍNDICE GENERAL
    # ══════════════════════════════════════════════════════════════
    doc.add_heading('ÍNDICE GENERAL', level=1)
    indice_items = [
        'PARTE I — DIAGRAMA ENTIDAD-RELACIÓN (ER)',
        '  1. Estructura General',
        '  2. Diagrama ER Completo',
        '  3. Lista de Tablas por Módulo',
        '  4. Diagramas Detallados por Módulo',
        '  5. Relaciones Clave',
        '  6. Restricciones y Validaciones',
        '  7. Particionamiento',
        '  8. Índices',
        '  9. Funciones SQL y Triggers',
        '  10. Tipos ENUM',
        '  11. Módulo M:M y Configuración',
        '  12. Módulo de Pedido Semanal a Bodega',
        '  13. Módulo de Solicitudes',
        '  14. Módulo de Pedidos',
        '  15. Sistema de Permisos',
        '',
        'PARTE II — ARQUITECTURA 4+1 VIEWS',
        '  1. Vista Lógica (Logical View)',
        '  2. Vista de Procesos (Process View)',
        '  3. Vista Física (Physical View)',
        '  4. Vista de Desarrollo (Development View)',
        '  5. Vista de Escenarios (Scenarios View)',
    ]
    for item in indice_items:
        if item == '':
            doc.add_paragraph()
        elif item.startswith('PARTE'):
            p = doc.add_paragraph(item)
            p.runs[0].bold = True
        else:
            doc.add_paragraph(item, style='List Number')

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════
    # PARTE I — DIAGRAMA ENTIDAD-RELACIÓN
    # ══════════════════════════════════════════════════════════════════════
    doc.add_heading('PARTE I — DIAGRAMA ENTIDAD-RELACIÓN (ER)', level=1)
    doc.add_paragraph('KuHub v1.0.8 — PostgreSQL 16.13')
    doc.add_paragraph()

    # ─── 1. ESTRUCTURA GENERAL ───
    doc.add_heading('1. Estructura General', level=2)
    doc.add_paragraph(
        'KuHub está compuesto por 9 módulos principales con 38 tablas totales:'
    )

    add_table_from_data(doc,
        ['Módulo', 'Tablas', 'Descripción'],
        [
            ['🔐 Seguridad', '5', 'usuario, rol, refresh_token, modulo, permiso_rol'],
            ['🎓 Académico', '7', 'asignatura, seccion, bloque_horario, sala, reserva_sala, docente_seccion, asignatura_profesor_cargo'],
            ['📦 Inventario', '6', 'unidad_medida, categoria, producto, inventario, bodega_transito, movimiento*'],
            ['🚚 Proveedores', '3', 'proveedor, proveedor_producto, proveedor_dia_entrega'],
            ['📦 Pedido Semanal a Bodega', '2', 'pedido_semana_bodega, detalle_pedido_semana_bodega'],
            ['🛒 Solicitudes', '3', 'solicitud, detalle_solicitud, motivo_rechazo_solicitud'],
            ['📝 Pedidos', '3', 'pedido, detalle_pedido, pedido_solicitud'],
            ['⚙️ Configuración', '2', 'gestion_sistema, semanas'],
            ['TOTAL', '38', ''],
        ]
    )
    doc.add_paragraph()

    # ─── 2. DIAGRAMA ER COMPLETO ───
    doc.add_heading('2. Diagrama ER Completo', level=2)
    doc.add_paragraph(
        'A continuación se presenta el diagrama ER completo del sistema KuHub. '
        'Las relaciones se representan con la notación estándar (1:N, M:M, 1:1).'
    )
    doc.add_paragraph()

    er_diagram = """
┌─────────────────────────────────────────────────────────────────────────────────────┐
│                           DIAGRAMA ER — KuHub v1.0.8                                │
├─────────────────────────────────────────────────────────────────────────────────────┤
│                                                                                     │
│  ┌─────────┐ 1:N  ┌──────────┐ 1:N  ┌──────────────┐                              │
│  │   ROL   │─────→│ USUARIO  │─────→│REFRESH_TOKEN │                              │
│  └─────────┘      └──────────┘      └──────────────┘                              │
│       │ 1:N            │ 1:N                                                        │
│       ↓                ├──────────────────────────────────┐                         │
│  ┌───────────┐         │                                  │                         │
│  │PERMISO_ROL│         ↓ 1:N                              ↓ 1:N                     │
│  └───────────┘    ┌────────────────┐              ┌─────────────┐                   │
│       ↑ 1:N       │DOCENTE_SECCION │              │ MOVIMIENTO* │                   │
│  ┌────────┐       └────────────────┘              └─────────────┘                   │
│  │ MODULO │              ↑ N:1                          ↑ N:1                       │
│  └────────┘              │                              │                           │
│                    ┌─────────────┐ 1:N  ┌────────────┐  │                           │
│  ┌────────────┐    │  SECCION    │─────→│RESERVA_SALA│  │                           │
│  │ASIGNATURA  │───→│             │      └────────────┘  │                           │
│  └────────────┘1:N └─────────────┘           ↑ N:1      │                           │
│       │ 1:N              │ 1:N          ┌────────┐      │                           │
│       ↓                  ↓              │  SALA  │      │                           │
│  ┌──────────────────┐  ┌──────────┐    └────────┘      │                           │
│  │ASIG_PROF_CARGO   │  │SOLICITUD │         ↑ N:1      │                           │
│  └──────────────────┘  └──────────┘    ┌────────────┐   │                           │
│                              │ 1:N     │BLOQUE_HORA │   │                           │
│                              ↓         └────────────┘   │                           │
│                    ┌──────────────────┐                  │                           │
│                    │DETALLE_SOLICITUD │                  │                           │
│                    └──────────────────┘                  │                           │
│                              ↑ N:1                      │                           │
│                              │                          │                           │
│  ┌───────────┐ 1:N  ┌───────────┐ 1:1  ┌───────────┐ 1:1 ┌──────────────┐         │
│  │ CATEGORIA │──────→│ PRODUCTO  │─────→│INVENTARIO │────→│BODEGA_TRANSITO│         │
│  └───────────┘       └───────────┘      └───────────┘     └──────────────┘         │
│                           ↑ N:1              │ 1:N              │ 1:N               │
│  ┌────────────┐           │                  └──────────────────┘                   │
│  │UNIDAD_MEDIDA│──────────┘ 1:N                     │                               │
│  └────────────┘           │                         ↓                               │
│                    ┌──────────────────┐      ┌─────────────┐                        │
│                    │PROVEEDOR_PRODUCTO│      │ MOVIMIENTO* │ (PARTICIONADO)          │
│                    └──────────────────┘      └─────────────┘                        │
│                           ↑ N:1                                                     │
│                    ┌───────────┐ 1:N  ┌─────────────────────┐                       │
│                    │ PROVEEDOR │─────→│PROVEEDOR_DIA_ENTREGA│                       │
│                    └───────────┘      └─────────────────────┘                       │
│                                                                                     │
│  ┌─────────┐ 1:N  ┌──────────────────────┐ 1:N  ┌─────────────────────────────┐   │
│  │ SEMANAS │─────→│PEDIDO_SEMANA_BODEGA  │─────→│DETALLE_PEDIDO_SEMANA_BODEGA │   │
│  └─────────┘      └──────────────────────┘      └─────────────────────────────┘   │
│                                                                                     │
│  ┌─────────┐ 1:N  ┌────────┐ M:N  ┌─────────────┐                                  │
│  │ PEDIDO  │─────→│DETALLE │─────→│ SOLICITUD   │                                  │
│  └─────────┘      └────────┘      └─────────────┘                                  │
│                                                                                     │
│  TOTALES: 9 MÓDULOS | 38 TABLAS | 6 RELACIONES M:M | 2 PARTICIONES | 47 ÍNDICES   │
│                                                                                     │
└─────────────────────────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, er_diagram)
    doc.add_paragraph()

    # ─── 3. LISTA DE TABLAS POR MÓDULO ───
    doc.add_heading('3. Lista de Tablas por Módulo', level=2)

    modulos_tablas = [
        ('🔐 MÓDULO SEGURIDAD (5 tablas)', [
            '• usuario (id, username, email, password_hash, activo, created_at)',
            '• rol (id, nombre, descripcion, activo)',
            '• permiso_rol (id, rol_id, modulo_id, puedeLeer, puedeCrear, puedeActualizar, puedeEliminar)',
            '• modulo (id, nombre, descripcion)',
            '• refresh_token (id, usuario_id, token, expira_en)',
        ]),
        ('🎓 MÓDULO ACADÉMICO (7 tablas)', [
            '• asignatura (id, codigo, nombre, descripcion, creditos)',
            '• seccion (id, asignatura_id, numero_seccion, semestre, año, cupos)',
            '• bloque_horario (id, seccion_id, dia, hora_inicio, hora_fin, sala_id)',
            '• sala (id, codigo, capacidad, ubicacion)',
            '• reserva_sala (id, sala_id, bloque_horario_id, fecha, estado)',
            '• docente_seccion (id, usuario_id, seccion_id, rol_docente)',
            '• asignatura_profesor_cargo (id, asignatura_id, usuario_id, cargo)',
        ]),
        ('📦 MÓDULO INVENTARIO (6 tablas)', [
            '• categoria (id, nombre, descripcion)',
            '• unidad_medida (id, nombre, abreviatura, conversion_factor)',
            '• producto (id, codigo, nombre, descripcion, categoria_id, unidad_medida_id, precio)',
            '• inventario (id, producto_id, stock_actual, stock_minimo, stock_maximo, ultimo_movimiento)',
            '• bodega_transito (id, producto_id, cantidad, fecha_llegada, estado)',
            '• movimiento (id, producto_id, tipo, cantidad, fecha, usuario_id, referencia) [PARTICIONADO POR MES]',
        ]),
        ('🚚 MÓDULO PROVEEDORES (3 tablas)', [
            '• proveedor (id, nombre, contacto, telefono, email, direccion)',
            '• proveedor_producto (id, proveedor_id, producto_id, precio_compra, tiempo_entrega)',
            '• proveedor_dia_entrega (id, proveedor_id, dia_semana, horario_inicio, horario_fin)',
        ]),
    ]

    for modulo_nombre, tablas in modulos_tablas:
        p = doc.add_paragraph(modulo_nombre)
        p.runs[0].bold = True
        for tabla in tablas:
            doc.add_paragraph(tabla, style='List Bullet')

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════
    # PARTE II — ARQUITECTURA 4+1 VIEWS
    # ══════════════════════════════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading('PARTE II — ARQUITECTURA 4+1 VIEWS', level=1)
    doc.add_paragraph('Arquitectura de Sistema para KuHub v1.0.8')
    doc.add_paragraph()

    # Vista Lógica
    doc.add_heading('1. Vista Lógica (Logical View)', level=2)
    doc.add_paragraph('Componentes principales y responsabilidades del sistema.')

    vista_logica = """
┌─────────────────────────────────────────────────────────────────────┐
│                      VISTA LÓGICA — KuHub                           │
├─────────────────────────────────────────────────────────────────────┤
│                                                                     │
│  FRONTEND (React 18 + TypeScript + Vite)                            │
│  ├─ Componentes UI                                                  │
│  │  ├─ LoginPage                                                   │
│  │  ├─ DashboardPage                                               │
│  │  ├─ SolicitudPage                                               │
│  │  ├─ InventarioPage                                              │
│  │  └─ GestionPage                                                 │
│  ├─ Services (API calls vía Axios)                                 │
│  │  ├─ AuthService (login, refresh token)                         │
│  │  ├─ SolicitudService (CRUD solicitudes)                         │
│  │  ├─ InventarioService (CRUD productos)                          │
│  │  └─ DashboardService (reportes y KPIs)                          │
│  └─ State Management (Context API)                                 │
│     ├─ AuthContext (JWT, usuario actual)                           │
│     └─ AppContext (datos globales)                                 │
│                                                                     │
│  BACKEND (Spring Boot 3 + Java 21)                                  │
│  ├─ Controllers                                                     │
│  │  ├─ AuthController (/login, /refresh)                          │
│  │  ├─ SolicitudController (CRUD solicitud)                        │
│  │  ├─ InventarioController (CRUD productos)                       │
│  │  ├─ DashboardController (reportes)                              │
│  │  └─ AdminController (gestión usuarios/roles)                    │
│  ├─ Services                                                        │
│  │  ├─ AuthService (validación JWT)                                │
│  │  ├─ SolicitudService (lógica negocio)                           │
│  │  ├─ InventarioService (stock, movimientos)                      │
│  │  └─ PermisoService (validación permisos)                        │
│  ├─ Repositories (JPA)                                              │
│  │  ├─ UsuarioRepository                                           │
│  │  ├─ SolicitudRepository                                         │
│  │  ├─ ProductoRepository                                          │
│  │  └─ MovimientoRepository                                        │
│  └─ Security                                                        │
│     ├─ JwtProvider (generación/validación token)                   │
│     ├─ SecurityConfig (Spring Security)                            │
│     └─ AuthenticationFilter (interceptor)                          │
│                                                                     │
│  DATABASE (PostgreSQL 16.13)                                        │
│  ├─ Schemas: kuhub (principal)                                      │
│  ├─ 38 Tablas                                                       │
│  ├─ 47 Índices optimizados                                          │
│  ├─ 6 Relaciones M:M                                                │
│  ├─ 2 Particiones (movimiento por mes)                              │
│  ├─ 8 Funciones SQL                                                 │
│  └─ 12 Triggers (validaciones)                                      │
│                                                                     │
└─────────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, vista_logica)
    doc.add_paragraph()

    # Vista de Procesos
    doc.add_heading('2. Vista de Procesos (Process View)', level=2)
    doc.add_paragraph('Flujos principales del sistema por rol.')

    vista_procesos = """
┌─────────────────────────────────────────────────────────────────────┐
│              VISTA DE PROCESOS — Flujos Principales                 │
├─────────────────────────────────────────────────────────────────────┤
│                                                                     │
│  FLUJO 1: Docente crea Solicitud de Productos                       │
│  ───────────────────────────────────────────────────────            │
│  1. Docente → ProtectedRoute: Valida token JWT                     │
│  2. ProtectedRoute → AuthService: Verifica permisos                │
│  3. AuthService → PermisoService: Matriz permiso_rol               │
│  4. ✅ Permitido → Carga SolicitudPage                            │
│  5. Docente → SelectProductos → Cantidad × Producto                │
│  6. Envía POST /api/solicitud/crear                                │
│  7. Backend → Validación de stock                                  │
│  8. DB → Inserta solicitud + detalles                              │
│  9. ✅ Response: ID_solicitud, estado=PENDIENTE                   │
│  10. Notificación → Profesor a Cargo                               │
│                                                                     │
│  FLUJO 2: Profesor a Cargo aprueba/rechaza Solicitud               │
│  ─────────────────────────────────────────────────────────          │
│  1. Profesor → /gestion-solicitudes                                 │
│  2. Carga lista de solicitudes PENDIENTES                           │
│  3. Selecciona solicitud → Revisa detalles                         │
│  4. ¿Aprueba? → Pulsa "Aceptar"                                    │
│     - Backend: UPDATE solicitud SET estado='ACEPTADA'              │
│     - Genera movimiento de inventario                              │
│  5. ¿Rechaza? → Pulsa "Rechazar" + motivo                         │
│     - Backend: INSERT motivo_rechazo_solicitud                     │
│     - UPDATE solicitud SET estado='RECHAZADA'                      │
│  6. Notificación → Docente (vía email)                             │
│                                                                     │
│  FLUJO 3: Encargado Bodega recibe productos                        │
│  ────────────────────────────────────────────                      │
│  1. Mensajero llega con pedido                                      │
│  2. Cajas → Zona Recepción                                          │
│  3. Sistema registra BodegaTransito (estado=PENDIENTE)              │
│  4. Enc. Bodega → /bodega-transito                                 │
│  5. Verifica cantidades vs Albarán                                 │
│  6. Pulsa "Recibir en Bodega"                                      │
│  7. Backend:                                                        │
│     - UPDATE inventario SET stock += cantidad                      │
│     - INSERT movimiento (tipo=INGRESO)                             │
│     - UPDATE bodega_transito SET estado='RECIBIDA'                 │
│  8. ✅ Producto disponible para solicitudes                        │
│                                                                     │
│  FLUJO 4: Gestor Pedidos crea Pedido a Proveedor                   │
│  ───────────────────────────────────────────────────                │
│  1. Gestor → /gestion-pedidos                                      │
│  2. Selecciona productos con bajo stock                            │
│  3. Valida: ¿Proveedor disponible? ¿Costo? ¿Tiempo?              │
│  4. Crea Pedido → PUT /api/pedido/crear                           │
│  5. Backend:                                                        │
│     - Inserta pedido + detalles                                    │
│     - Valida disponibilidad de presupuesto                         │
│  6. Estado: CREADO → Espera confirmación proveedor                 │
│                                                                     │
│  FLUJO 5: Admin configura Permisos                                 │
│  ─────────────────────────────                                     │
│  1. Admin → /gestion-roles                                         │
│  2. Selecciona Rol → Módulos y permisos                            │
│  3. Modifica: puedeLeer, puedeCrear, puedeActualizar               │
│  4. Guarda cambios → POST /api/rol/actualizar                      │
│  5. Backend: UPDATE permiso_rol                                    │
│  6. Cache invalidado → Siguiente login usa permisos nuevos         │
│                                                                     │
└─────────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, vista_procesos)
    doc.add_paragraph()

    # Vista Física
    doc.add_heading('3. Vista Física (Physical View)', level=2)
    doc.add_paragraph('Infraestructura de despliegue en AWS.')

    vista_fisica = """
┌─────────────────────────────────────────────────────────────────────┐
│           VISTA FÍSICA — Infraestructura AWS Lightsail              │
├─────────────────────────────────────────────────────────────────────┤
│                                                                     │
│  REGIÓN: us-east-1                                                  │
│                                                                     │
│  ┌──────────────────────────────────────────────────────────┐      │
│  │                   INTERNET / DNS                         │      │
│  │                  (Route 53)                              │      │
│  └────────────────┬─────────────────────────────────────────┘      │
│                   │                                                 │
│                   ↓                                                 │
│  ┌──────────────────────────────────────────────────────────┐      │
│  │          AWS LIGHTSAIL LOAD BALANCER                     │      │
│  │          (HTTPS:443 → Backend:8080, Frontend)            │      │
│  └────────────┬─────────────────────────────────────────────┘      │
│               │                                                     │
│       ┌───────┴──────────┐                                         │
│       ↓                  ↓                                         │
│  ┌─────────────┐   ┌─────────────────────┐                        │
│  │  FRONTEND   │   │  BACKEND INSTANCE   │                        │
│  │ INSTANCE 1  │   │  (Spring Boot 3)    │                        │
│  ├─────────────┤   ├─────────────────────┤                        │
│  │ React 18    │   │ Java 21             │                        │
│  │ TypeScript  │   │ Port: 8080          │                        │
│  │ Vite        │   │ Services/Repos      │                        │
│  │ Nginx       │   │ JPA, Hibernate      │                        │
│  │ Cert: SSL   │   │ Auth/Security       │                        │
│  └─────────────┘   └──────────┬──────────┘                        │
│                               │                                    │
│                      ┌────────┴──────────┐                        │
│                      │                   │                        │
│                      ↓                   ↓                        │
│               ┌──────────────┐    ┌──────────────┐               │
│               │  DB INSTANCE │    │  DB STANDBY  │               │
│               │ PostgreSQL   │    │  (Replica)   │               │
│               │   16.13      │    │  Read-only   │               │
│               │ 38 Tablas    │    │              │               │
│               │ 47 Índices   │    │              │               │
│               │ 2 Particiones│    │              │               │
│               └──────────────┘    └──────────────┘               │
│                      │                   │                        │
│                      └────────┬──────────┘                        │
│                               ↓                                   │
│                  ┌────────────────────────┐                       │
│                  │   BACKUP S3 BUCKET     │                       │
│                  │   (Automatizado diario) │                       │
│                  └────────────────────────┘                       │
│                                                                     │
│  SEGURIDAD:                                                         │
│  • VPC privada para DB                                             │
│  • Security Groups (restrict by port/protocol)                     │
│  • SSL/TLS para conexiones                                         │
│  • Snapshots automáticos cada 6 horas                              │
│                                                                     │
└─────────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, vista_fisica)
    doc.add_paragraph()

    # Vista de Desarrollo
    doc.add_heading('4. Vista de Desarrollo (Development View)', level=2)
    doc.add_paragraph('Estructura del código fuente y módulos.')

    vista_desarrollo = """
┌─────────────────────────────────────────────────────────────────────┐
│           VISTA DE DESARROLLO — Estructura de Código                │
├─────────────────────────────────────────────────────────────────────┤
│                                                                     │
│  BACKEND (Spring Boot)                                              │
│  src/main/java/KuHub/                                               │
│  ├─ config/                  (Configuración Spring)                │
│  │  ├─ SecurityConfig.java                                        │
│  │  ├─ JwtConfig.java                                             │
│  │  └─ DatabaseConfig.java                                        │
│  ├─ security/                (Autenticación y Autorización)        │
│  │  ├─ JwtProvider.java                                           │
│  │  ├─ SecurityFilter.java                                        │
│  │  └─ PermisoValidator.java                                      │
│  ├─ modules/                 (Módulos de negocio)                 │
│  │  ├─ gestion_usuario/      (Usuarios, Roles, Permisos)          │
│  │  ├─ modulo_academico/     (Asignaturas, Secciones, Salas)      │
│  │  ├─ modulo_inventario/    (Productos, Stock, Movimientos)      │
│  │  ├─ modulo_solicitud/     (Solicitudes, Aprobaciones)          │
│  │  ├─ modulo_pedido/        (Pedidos a Proveedores)              │
│  │  └─ modulo_configuracion/ (Sistemas generales)                 │
│  ├─ controllers/             (REST Endpoints)                     │
│  ├─ services/                (Lógica negocio)                     │
│  ├─ repositories/            (JPA Data Access)                    │
│  ├─ dto/                     (Data Transfer Objects)              │
│  ├─ models/                  (Entity JPA)                         │
│  └─ exceptions/              (Manejo errores)                     │
│                                                                     │
│  FRONTEND (React)                                                   │
│  src/                                                               │
│  ├─ components/              (Componentes reutilizables)          │
│  │  ├─ Dashboard/                                                 │
│  │  ├─ Solicitud/                                                 │
│  │  ├─ Inventario/                                                │
│  │  ├─ Gestion/                                                   │
│  │  └─ Auth/                                                      │
│  ├─ pages/                   (Páginas principales)                │
│  │  ├─ LoginPage.tsx                                              │
│  │  ├─ DashboardPage.tsx                                          │
│  │  ├─ SolicitudPage.tsx                                          │
│  │  └─ GestionPage.tsx                                            │
│  ├─ services/                (API calls)                          │
│  │  ├─ AuthService.ts                                             │
│  │  ├─ SolicitudService.ts                                        │
│  │  ├─ InventarioService.ts                                       │
│  │  └─ DashboardService.ts                                        │
│  ├─ context/                 (Estado global)                      │
│  │  ├─ AuthContext.ts                                             │
│  │  └─ AppContext.ts                                              │
│  ├─ hooks/                   (Custom React hooks)                 │
│  ├─ types/                   (TypeScript interfaces)              │
│  ├─ utils/                   (Funciones auxiliares)               │
│  └─ styles/                  (CSS/TailwindCSS)                    │
│                                                                     │
│  DATABASE (PostgreSQL)                                              │
│  ├─ schema.sql               (DDL: CREATE TABLE)                  │
│  ├─ functions.sql            (Funciones SQL)                      │
│  ├─ triggers.sql             (Triggers para validaciones)         │
│  ├─ indexes.sql              (Índices optimizados)                │
│  ├─ partitions.sql           (Particionamiento)                   │
│  └─ seed.sql                 (Datos iniciales)                    │
│                                                                     │
└─────────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, vista_desarrollo)
    doc.add_paragraph()

    # Vista de Escenarios
    doc.add_heading('5. Vista de Escenarios (Scenarios View)', level=2)
    doc.add_paragraph('Casos de uso principales por rol.')

    # 5.1 Docente solicita productos
    doc.add_heading('5.1 Caso de Uso: Docente solicita productos para clase', level=3)
    doc.add_paragraph(
        'Rol: Docente (ID: 5)\n'
        'Página: /solicitud\n'
        'Permisos: SOLICITUD.puedeCrear = true'
    )
    caso_uso_docente = """
┌──────────────────────────────────────────────────────────────────────────────┐
│ CASO DE USO: DOCENTE SOLICITA PRODUCTOS PARA CLASE                          │
├──────────────────────────────────────────────────────────────────────────────┤
│                                                                              │
│  1. Docente inicia sesión → /login                                           │
│  2. Ingresa credenciales (usuario: dgarcia, password: **)                   │
│  3. Backend → AuthController /login                                         │
│     - Valida usuario existe en DB                                           │
│     - Verifica password_hash                                                │
│     - Genera JWT (válido 1 hora)                                            │
│  4. ✅ Response: { token, usuario: { id, nombre, rol_id } }                │
│  5. Frontend → LocalStorage: guarda JWT                                     │
│  6. Navega a /solicitud                                                     │
│  7. ProtectedRoute → Verifica JWT en header Authorization                   │
│  8. JwtProvider → Valida firma y expiration                                 │
│  9. PermisoService → Consulta permiso_rol:                                 │
│     SELECT * FROM permiso_rol WHERE rol_id=5 AND modulo_id='SOLICITUD'     │
│  10. ✅ puedeLeer=true, puedeCrear=true → Renderiza SolicitudPage         │
│  11. SolicitudService → GET /api/solicitud/mis-solicitudes                  │
│  12. Backend → Retorna solicitudes where usuario_id=5                       │
│  13. Docente ve tabla: [Asignatura | Sección | Estado | Fecha]             │
│  14. Pulsa "Nueva Solicitud"                                                │
│  15. Modal abre → SelectProducto (dropdown de inventario)                   │
│  16. Docente selecciona:                                                    │
│      - Producto: "Papel A4"                                                 │
│      - Cantidad: 5 resmas                                                   │
│  17. Agrega fila → Repite para más productos                                │
│  18. Confirma → POST /api/solicitud/crear                                  │
│      Body: { usuario_id: 5, detalles: [{producto_id, cantidad}] }          │
│  19. Backend → SolicitudService.crear():                                    │
│      - Valida stock disponible                                              │
│      - INSERT solicitud (estado='PENDIENTE')                               │
│      - INSERT detalle_solicitud (x por cada producto)                       │
│      - Genera notificación → Profesor a Cargo                               │
│  20. ✅ Response: { solicitud_id: 42, estado: 'PENDIENTE' }               │
│  21. Frontend → Toast: "Solicitud creada exitosamente"                      │
│  22. Tabla actualiza → Nueva solicitud aparece con estado PENDIENTE         │
│  23. Docente espera respuesta del Profesor a Cargo                          │
│                                                                              │
│  MATRIZ DE PERMISOS — DOCENTE EN MÓDULO SOLICITUD:                          │
│  ┌──────────────────────┬──────────────┬──────────────────────────────┐     │
│  │ Acción               │ Disponible   │ Detalles                     │     │
│  ├──────────────────────┼──────────────┼──────────────────────────────┤     │
│  │ Ver solicitudes      │ ✅ Sí        │ puedeLeer = TRUE             │     │
│  │ Crear solicitud      │ ✅ Sí        │ puedeCrear = TRUE            │     │
│  │ Editar solicitud     │ ❌ No        │ puedeActualizar = FALSE      │     │
│  │ Eliminar solicitud   │ ❌ No        │ puedeEliminar = FALSE        │     │
│  └──────────────────────┴──────────────┴──────────────────────────────┘     │
│                                                                              │
└──────────────────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, caso_uso_docente)
    doc.add_paragraph()

    # 5.2 Profesor a Cargo aprueba/rechaza
    doc.add_heading('5.2 Caso de Uso: Profesor a Cargo aprueba solicitud', level=3)
    doc.add_paragraph(
        'Rol: Profesor a Cargo (ID: 4)\n'
        'Página: /gestion-solicitudes\n'
        'Permisos: SOLICITUD.puedeLeer y SOLICITUD.puedeActualizar'
    )
    caso_uso_profesor = """
┌──────────────────────────────────────────────────────────────────────────────┐
│ CASO DE USO: PROFESOR A CARGO APRUEBA SOLICITUD                             │
├──────────────────────────────────────────────────────────────────────────────┤
│                                                                              │
│  1. Profesor a Cargo inicia sesión                                           │
│  2. Navega a /gestion-solicitudes                                            │
│  3. ProtectedRoute:                                                          │
│     - Valida JWT                                                             │
│     - pageId = 'gestion-solicitudes'                                         │
│     - Consulta matriz: SELECT pageId FROM usuario_rol WHERE rol_id=4        │
│  4. ✅ Acceso permitido                                                      │
│  5. SolicitudService → GET /api/solicitud/pendientes                        │
│     SELECT * FROM solicitud WHERE estado='PENDIENTE' AND created_by_id=?   │
│  6. Tabla carga: [Docente | Asignatura | Productos | Fecha | Acciones]    │
│  7. Profesor selecciona solicitud del Docente García                        │
│  8. ¿Aprueba? → Pulsa "Aceptar"                                            │
│     Backend → SolicitudService.aceptar(solicitud_id):                       │
│     - UPDATE solicitud SET estado='ACEPTADA', fecha_aprobacion=NOW()       │
│     - INSERT movimiento (tipo=DESPACHO, ref=solicitud_id)                  │
│     - Actualiza inventario: stock -= cantidad                              │
│     - Genera PDF de solicitud aprobada                                      │
│     - Envía email a Docente                                                 │
│  9. ¿Rechaza? → Pulsa "Rechazar"                                           │
│     Modal abre → Selecciona motivo:                                         │
│     - Sin stock disponible                                                  │
│     - Producto discontinuado                                                │
│     - Fuera de presupuesto                                                  │
│     Backend → SolicitudService.rechazar(solicitud_id, motivo):             │
│     - INSERT motivo_rechazo_solicitud                                       │
│     - UPDATE solicitud SET estado='RECHAZADA'                               │
│     - Envía email a Docente con motivo                                      │
│  10. ✅ Solicitud procesada                                                 │
│  11. Tabla se actualiza → Solicitud ya no aparece en PENDIENTES             │
│                                                                              │
└──────────────────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, caso_uso_profesor)
    doc.add_paragraph()

    # 5.3 Gestor de Pedidos
    doc.add_heading('5.3 Caso de Uso: Gestor de Pedidos crea compra', level=3)
    doc.add_paragraph(
        'Rol: Gestor de Pedidos (ID: 8)\n'
        'Página: /gestion-pedidos\n'
        'Permisos: PEDIDO.puedeLeer, PEDIDO.puedeCrear'
    )
    caso_uso_gestor = """
┌──────────────────────────────────────────────────────────────────────────────┐
│ CASO DE USO: GESTOR DE PEDIDOS CREA COMPRA A PROVEEDOR                      │
├──────────────────────────────────────────────────────────────────────────────┤
│                                                                              │
│  1. Gestor de Pedidos → /gestion-pedidos                                    │
│  2. ProtectedRoute: pageId='gestion-pedidos', Valida permisos               │
│  3. ✅ Permitido                                                            │
│  4. Dashboard:                                                              │
│     - Tab 1: Productos con stock bajo (<stock_minimo)                       │
│     - Tab 2: Proveedores disponibles                                        │
│     - Tab 3: Pedidos pendientes                                             │
│  5. Gestor revisa productos críticos:                                       │
│     SELECT * FROM inventario WHERE stock < stock_minimo ORDER BY stock      │
│     Resultado:                                                              │
│     - Papel A4: 2 resmas (mín: 10)                                         │
│     - Tóner HP: 1 cartuchos (mín: 5)                                        │
│  6. Gestor selecciona "Papel A4" → Crea nuevo Pedido                        │
│  7. Modal abre → Selecciona:                                                │
│     - Proveedor: "Distribuidora Central"                                    │
│     - Cantidad: 20 resmas                                                   │
│     - Validaciones:                                                         │
│       * Consulta proveedor_producto:                                        │
│         SELECT * FROM proveedor_producto                                    │
│         WHERE proveedor_id=X AND producto_id=Y                              │
│       * Obtiene: precio_compra=500, tiempo_entrega=2 días                  │
│  8. Calcula: Costo total = 20 × 500 = 10,000 CLP                           │
│  9. Valida presupuesto → Consulta gestion_sistema.presupuesto_disponible   │
│  10. ✅ Presupuesto suficiente                                             │
│  11. Confirma → POST /api/pedido/crear                                      │
│      Body: { proveedor_id, detalles: [{producto_id, cantidad}], costo }   │
│  12. Backend → PedidoService.crear():                                       │
│      - INSERT pedido (estado='CREADO')                                      │
│      - INSERT detalle_pedido                                                │
│      - UPDATE gestion_sistema SET presupuesto_disponible -= 10,000          │
│      - Genera número de OC (purchase_order_number)                          │
│      - Envía email a Proveedor con OC                                       │
│  13. ✅ Response: { pedido_id, estado: 'CREADO', numero_oc: 'OC-2026-042' }│
│  14. Gestor puede revisar historial de pedidos                              │
│  15. Estado esperado: CREADO → CONFIRMADO → ENTREGADO                       │
│                                                                              │
└──────────────────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, caso_uso_gestor)
    doc.add_paragraph()

    # 5.4 Encargado de Bodega
    doc.add_heading('5.4 Caso de Uso: Encargado bodega recibe y registra productos', level=3)
    doc.add_paragraph(
        'Rol: Encargado de Bodega (ID: 6) o Asistente (ID: 7)\n'
        'Página: /bodega-transito\n'
        'Permisos: BODEGA_TRANSITO.puedeActualizar'
    )
    caso_uso_bodega = """
┌──────────────────────────────────────────────────────────────────────────────┐
│ CASO DE USO: ENCARGADO BODEGA RECIBE Y REGISTRA PRODUCTOS                   │
├──────────────────────────────────────────────────────────────────────────────┤
│                                                                              │
│  1. Mensajero llega con pedido confirmado                                    │
│  2. Cajas se colocan en Zona de Recepción                                    │
│  3. Sistema registra items en BodegaTransito (estado: PENDIENTE)             │
│  4. Enc. Bodega accede a /bodega-transito                                    │
│     Permiso: BODEGA_TRANSITO.puedeLeer                                       │
│  5. Sistema carga lista de items pendientes                                  │
│  6. Abre cada caja, verifica cantidades contra albarán/factura               │
│  7. Revisa estado físico: ¿Rotos? ¿Vencidos? ¿Contaminados?                 │
│  8. Pulsa "Recibir en Bodega"                                                │
│     Permiso: BODEGA_TRANSITO.puedeActualizar                                 │
│  9. Modal: Confirma cantidad recibida (puede ajustar)                        │
│  10. Sistema ejecuta:                                                        │
│      - UPDATE inventario SET stock += cantidadReal                            │
│      - INSERT movimiento (tipo: INGRESO, referencia: tránsito)               │
│      - UPDATE bodega_transito SET estado = 'RECIBIDA'                        │
│  11. ✅ Producto ingresado al inventario activo                              │
│  12. Disponible para solicitudes de docentes                                 │
│                                                                              │
│  CAMBIOS DE ESTADO:                                                          │
│  BodegaTransito: PENDIENTE → RECIBIDO (confirmado, pasa a Inventario)        │
│                                                                              │
│  PERMISOS POR ROL:                                                           │
│  ┌──────────────────────┬────────┬────────┬────────┬──────────────┐          │
│  │ Rol                  │ Acceso │ Crear  │ Leer   │ Actualizar   │          │
│  ├──────────────────────┼────────┼────────┼────────┼──────────────┤          │
│  │ Encargado de Bodega  │ ✅ Sí  │ ❌ No  │ ✅ Sí  │ ✅ Sí        │          │
│  │ Asistente de Bodega  │ ✅ Sí  │ ❌ No  │ ✅ Sí  │ ✅ Sí        │          │
│  │ Otros roles          │ ❌ No  │ ❌ No  │ ❌ No  │ ❌ No        │          │
│  └──────────────────────┴────────┴────────┴────────┴──────────────┘          │
│                                                                              │
└──────────────────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, caso_uso_bodega)
    doc.add_paragraph()

    # 5.5 Docente consulta
    doc.add_heading('5.5 Caso de Uso: Docente consulta estado de solicitudes', level=3)
    doc.add_paragraph(
        'Rol: Docente (ID: 5)\n'
        'Página: /solicitud\n'
        'Permisos: SOLICITUD.puedeLeer = true (lectura solamente)\n'
        'Diferencia con Profesor a Cargo: El Docente solo PUEDE LEER, NO puede crear.'
    )
    caso_uso_docente_consulta = """
┌──────────────────────────────────────────────────────────────────────────────┐
│ CASO DE USO: DOCENTE CONSULTA ESTADO DE SOLICITUDES                          │
├──────────────────────────────────────────────────────────────────────────────┤
│                                                                              │
│  1. Docente inicia sesión → JWT generado                                     │
│  2. ProtectedRoute verifica: SOLICITUD.puedeLeer=TRUE                        │
│  3. ✅ Acceso permitido PERO sin puedeCrear                                  │
│  4. Sistema carga solicitudes de sesiones del docente                        │
│  5. Filtra por estado: PENDIENTE, ACEPTADA, RECHAZADA                        │
│  6. Revisa detalles: Asignatura, Sección, Productos, Estado, Comentarios     │
│  7. Puede exportar reporte de solicitudes                                    │
│  8. Botón "Crear Solicitud" está DESHABILITADO (puedeCrear=FALSE)            │
│  9. Si necesita crear → solicita a Profesor a Cargo                          │
│                                                                              │
│  MATRIZ DE PERMISOS — DOCENTE EN MÓDULO SOLICITUD:                           │
│  ┌──────────────────────┬──────────────┬──────────────────────────────┐      │
│  │ Acción               │ Disponible   │ Detalles                     │      │
│  ├──────────────────────┼──────────────┼──────────────────────────────┤      │
│  │ Ver solicitudes      │ ✅ Sí        │ puedeLeer = TRUE             │      │
│  │ Crear solicitud      │ ❌ No        │ puedeCrear = FALSE           │      │
│  │ Editar solicitud     │ ❌ No        │ puedeActualizar = FALSE      │      │
│  │ Eliminar solicitud   │ ❌ No        │ puedeEliminar = FALSE        │      │
│  │ Exportar reporte     │ ✅ Sí        │ Función de lectura           │      │
│  └──────────────────────┴──────────────┴──────────────────────────────┘      │
│                                                                              │
└──────────────────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, caso_uso_docente_consulta)
    doc.add_paragraph()

    # 5.6 Asistente Bodega
    doc.add_heading('5.6 Caso de Uso: Asistente de Bodega apoya en recepción', level=3)
    caso_uso_asistente = """
┌──────────────────────────────────────────────────────────────────────────────┐
│ CASO DE USO: ASISTENTE DE BODEGA APOYA EN RECEPCIÓN                          │
├──────────────────────────────────────────────────────────────────────────────┤
│                                                                              │
│  1. Asistente inicia sesión                                                  │
│  2. Navega a /bodega-transito                                                │
│     ProtectedRoute verifica: pageId='bodega-transito'                        │
│  3. ✅ Acceso permitido                                                      │
│  4. Sistema carga lista de productos pendientes de confirmar                 │
│  5. Recibe instrucciones del Encargado de Bodega                             │
│  6. Abre cajas de recepción                                                  │
│  7. Verifica cantidades básicas con albarán                                  │
│  8. Puede registrar estado físico inicial                                    │
│  9. Pulsa "Registrar Recepción"                                              │
│     Backend verifica: BODEGA_TRANSITO.puedeCrear                             │
│  10. Completa datos: cantidad recibida, observaciones                        │
│  11. Guarda cambios                                                          │
│  12. Encargado revisa y confirma                                             │
│  13. ✅ Producto registrado, aguardando confirmación final                   │
│                                                                              │
│  RESPONSABILIDADES DEL ASISTENTE:                                            │
│  ✅ Ayudar en verificación física de productos                               │
│  ✅ Registrar cantidades iniciales                                           │
│  ✅ Documentar daños o discrepancias                                         │
│  ✅ Crear registros de movimientos                                           │
│  ❌ NO puede confirmar recepción final (solo Encargado)                      │
│  ❌ NO puede eliminar registros                                              │
│                                                                              │
└──────────────────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, caso_uso_asistente)
    doc.add_paragraph()

    # 5.7 Co-Admin
    doc.add_heading('5.7 Caso de Uso: Co-Administrador supervisa operaciones', level=3)
    caso_uso_coadmin = """
┌──────────────────────────────────────────────────────────────────────────────┐
│ CASO DE USO: CO-ADMINISTRADOR SUPERVISA OPERACIONES                          │
├──────────────────────────────────────────────────────────────────────────────┤
│                                                                              │
│  1. Co-Admin inicia sesión                                                   │
│  2. Navega a /dashboard                                                      │
│     ProtectedRoute verifica: Acceso a módulos operativos EXCEPTO Admin       │
│  3. ✅ Acceso permitido                                                      │
│  4. Visualiza dashboard con 3 tabs principales:                              │
│     📊 Tab 1: Inventario (Stock por categoría, Productos críticos)           │
│     📊 Tab 2: Solicitudes (Pendientes vs Aceptadas vs Rechazadas)            │
│     📊 Tab 3: Pedidos (Consolidados, Costo estimado)                         │
│  5. Puede navegar a:                                                         │
│     ✅ /gestion-solicitudes → Revisar y cambiar estado                       │
│     ✅ /gestion-pedidos → Crear pedidos consolidados                         │
│     ✅ /inventario → Ver y editar productos                                  │
│  6. ❌ RESTRICCIONES:                                                        │
│     NO: /gestion-roles                                                       │
│     NO: /gestion-usuarios                                                    │
│     NO: /admin-sistema                                                       │
│  7. 🎯 Resultado: Supervisión operativa completa                             │
│                                                                              │
│  COMPARATIVA ADMIN vs CO-ADMIN:                                              │
│  ┌──────────────────────────┬──────────────────┬──────────────────┐          │
│  │ Aspecto                  │ Administrador    │ Co-Administrador │          │
│  ├──────────────────────────┼──────────────────┼──────────────────┤          │
│  │ Gestión de Solicitudes   │ ✅ Sí            │ ✅ Sí            │          │
│  │ Gestión de Pedidos       │ ✅ Sí            │ ✅ Sí            │          │
│  │ Inventario               │ ✅ CRUD completo │ ✅ Crear/Editar  │          │
│  │ Gestión de Roles         │ ✅ Sí            │ ❌ NO            │          │
│  │ Gestión de Usuarios      │ ✅ Sí            │ ❌ NO            │          │
│  │ Admin del Sistema        │ ✅ Sí            │ ❌ NO            │          │
│  └──────────────────────────┴──────────────────┴──────────────────┘          │
│                                                                              │
└──────────────────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, caso_uso_coadmin)
    doc.add_paragraph()

    # 5.8 Admin Dashboard
    doc.add_heading('5.8 Caso de Uso: Administrador visualiza dashboards y reportes', level=3)
    caso_uso_admin_dash = """
┌──────────────────────────────────────────────────────────────────────────────┐
│ CASO DE USO: ADMINISTRADOR VISUALIZA DASHBOARDS Y REPORTES                   │
├──────────────────────────────────────────────────────────────────────────────┤
│                                                                              │
│  1. Admin inicia sesión → Accede a /dashboard                                │
│  2. Sistema detecta rol ADMINISTRADOR                                        │
│  3. Backend calcula tabs disponibles (4 tabs para Admin)                     │
│  4. Tabs disponibles:                                                        │
│     📊 Tab 1: Inventario                                                     │
│        - Gráfico: Stock por categoría                                        │
│        - Tabla: Productos críticos (stock bajo)                              │
│        - Histórico: Movimientos últimos 7 días                               │
│     📊 Tab 2: Solicitudes                                                    │
│        - KPI: Pendientes vs Aceptadas vs Rechazadas                          │
│        - Tabla: Solicitudes por semana                                       │
│        - Filtro: Por estado, asignatura                                      │
│     📊 Tab 3: Pedidos                                                        │
│        - KPI: Pedidos por semana                                             │
│        - Costo total estimado                                                │
│        - Ingredientes más solicitados                                        │
│     📊 Tab 4: Administración (Solo Admin)                                    │
│        - Usuarios activos/inactivos                                          │
│        - Cambios recientes                                                   │
│        - Auditoría de permisos                                               │
│  5. Admin interactúa: Selecciona filtros (fechas, asignatura, estado)        │
│  6. Sistema recalcula en tiempo real                                         │
│  7. Puede exportar reporte (PDF/Excel)                                       │
│  8. 🎉 Datos disponibles para decisiones estratégicas                        │
│                                                                              │
│  DIFERENCIAS EN DASHBOARDS POR ROL:                                          │
│  ┌──────────────────┬──────────┬──────────┬──────────┬──────────┐            │
│  │ Aspecto          │ Admin    │ Co-Admin │ Gestor   │ Bodega   │            │
│  ├──────────────────┼──────────┼──────────┼──────────┼──────────┤            │
│  │ Tab Inventario   │ ✅ Compl.│ ✅ Compl.│ ❌ No    │ ✅ Básico│            │
│  │ Tab Solicitudes  │ ✅ Compl.│ ✅ Compl.│ ✅ Compl.│ ❌ No    │            │
│  │ Tab Pedidos      │ ✅ Compl.│ ✅ Compl.│ ✅ Compl.│ ❌ No    │            │
│  │ Tab Administrac. │ ✅ Sí    │ ❌ No    │ ❌ No    │ ❌ No    │            │
│  │ Exportar report. │ ✅ PDF/XL│ ✅ PDF/XL│ ✅ PDF/XL│ ✅ PDF/XL│            │
│  └──────────────────┴──────────┴──────────┴──────────┴──────────┘            │
│                                                                              │
└──────────────────────────────────────────────────────────────────────────────┘
"""
    add_code_block(doc, caso_uso_admin_dash)
    doc.add_paragraph()

    # 5.9 Resumen de Restricciones
    doc.add_heading('5.9 Resumen de Restricciones por Rol', level=3)
    doc.add_paragraph().add_run('🔴 Restricciones Comunes:').bold = True
    add_table_from_data(doc,
        ['Restricción', 'Detalles'],
        [
            ['Acceso vía ProtectedRoute', 'Cada página valida pageId contra matriz de permisos. Si denegado → /sin-acceso'],
            ['Validación backend', 'TODO endpoint requiere JWT válido en header Authorization: Bearer <token>'],
            ['Control de módulos', 'Tabla permiso_rol mapea cada rol a módulos. Permisos CRUD son booleanos'],
            ['Eliminación lógica', 'Todos los datos se marcan activo=false, nunca se eliminan permanentemente'],
            ['Auditoría', 'Cada cambio registra usuario, timestamp y tipo de operación'],
        ]
    )
    doc.add_paragraph()

    # 5.10 Gestión Dinámica de Permisos
    doc.add_heading('5.10 Gestión Dinámica de Permisos', level=3)
    doc.add_paragraph(
        'El sistema permite a los Administradores cambiar los permisos en tiempo real '
        'desde /gestion-roles. Cualquier combinación de: puedeLeer, puedeCrear, '
        'puedeActualizar, puedeEliminar puede ser modificada.'
    )
    doc.add_paragraph()
    doc.add_paragraph(
        'Para restaurar permisos predeterminados:\n'
        '• Via UI: /gestion-roles → Seleccionar rol → "Restaurar Predeterminados" → Guardar\n'
        '• Via SQL: Ejecutar bloque de permisos en ConexionXD_v2.sql (líneas 825-1096)'
    )
    doc.add_paragraph()

    # Acceso Permitido por Página
    doc.add_paragraph().add_run('🟢 Acceso Permitido por Página:').bold = True
    acceso_pagina = """
✅ DASHBOARD
  ├─ Administrador (todos los tabs)
  ├─ Co-Administrador (sin tab Administración)
  ├─ Gestor de Pedidos (solo KPIs de pedidos)
  ├─ Profesor/Docente (resumen general)
  ├─ Encargado Bodega (KPIs de inventario)
  └─ Asistente Bodega (información general)

✅ SOLICITUD (crear)
  ├─ Docente (solo lectura)
  └─ Profesor a Cargo (crear/editar)

✅ GESTION_SOLICITUDES (revisar/aprobar)
  ├─ Administrador
  ├─ Co-Administrador
  └─ Gestor de Pedidos

✅ GESTION_PEDIDOS (crear/consolidar)
  ├─ Administrador
  ├─ Co-Administrador
  └─ Gestor de Pedidos

✅ INVENTARIO (consultar/editar)
  ├─ Administrador
  ├─ Co-Administrador
  └─ Encargado Bodega

✅ BODEGA_TRANSITO (recibir productos)
  ├─ Administrador
  ├─ Co-Administrador
  ├─ Encargado Bodega
  └─ Asistente Bodega

✅ GESTION_ROLES (administrar permisos)
  └─ Administrador (solo)

✅ GESTION_USUARIOS (crear/editar usuarios)
  └─ Administrador (solo)

✅ ADMIN_SISTEMA (configuración global)
  └─ Administrador (solo)
"""
    add_code_block(doc, acceso_pagina)
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # RESUMEN FINAL
    # ══════════════════════════════════════════════════════════════
    doc.add_page_break()
    doc.add_heading('RESUMEN COMPARATIVO DE VISTAS', level=1)
    add_table_from_data(doc,
        ['Vista', 'Foco', 'Audiencia', 'Cambio frecuencia'],
        [
            ['Lógica', 'Componentes y responsabilidades', 'Arquitectos, devs', 'Medio'],
            ['Procesos', 'Flujos y interacciones', 'PO, QA, backend devs', 'Bajo'],
            ['Física', 'Infraestructura y despliegue', 'DevOps, SRE, infra', 'Bajo'],
            ['Desarrollo', 'Estructura del código', 'Developers', 'Alto'],
            ['Escenarios', 'Casos de uso principales', 'Stakeholders, QA, PO', 'Bajo'],
        ]
    )
    doc.add_paragraph()

    # Pie de documento
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        '─────────────────────────────────────────────────────────────────\n'
        'Documento generado: 12 de mayo de 2026\n'
        'Sistema: KuHub v1.0.8\n'
        'Base de Datos: PostgreSQL 16.13 — 38 tablas, 9 módulos\n'
        'Backend: Spring Boot 3 + Java 21\n'
        'Frontend: React 18 + TypeScript + Vite\n'
        'Infraestructura: AWS Lightsail (2 instancias)\n'
        '─────────────────────────────────────────────────────────────────'
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(128, 128, 128)

    return doc


if __name__ == '__main__':
    print("Generando EVALUACION_PARCIAL_2_TPY1101.docx...")
    document = create_document()
    output_path = 'EVALUACION_PARCIAL_2_TPY1101.docx'
    document.save(output_path)
    print(f"[OK] Documento generado exitosamente: {output_path}")
    print(f"     Tamaño: {os.path.getsize(output_path) / 1024:.1f} KB")
